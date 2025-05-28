import os
import time
import json
import shutil
import streamlit as st
from pathlib import Path
from llama_index.core import (
    SimpleDirectoryReader,
    VectorStoreIndex,
    StorageContext,
    load_index_from_storage,
    Settings,
    Document
)
from llama_index.llms.openai import OpenAI
from llama_index.embeddings.openai import OpenAIEmbedding
import openai

# Configure page
st.set_page_config(
    page_title="📚 AI Teaching Assistant",
    layout="wide",
    page_icon="📚",
    initial_sidebar_state="expanded"
)

# Title and description
st.title("📚 AI Teaching Assistant")
st.markdown("""
Upload course materials to create an AI teaching assistant that can:
- Answer questions about the content
- Generate practice test questions
- Extract key concepts from lectures
""")

# Subject-specific configurations
SUBJECT_CONFIGS = {
    "economics": {
        "description": "Economic theories, models, and real-world applications",
        "guidance": """
        - Always reference economic principles like supply/demand, elasticity, opportunity cost
        - Use appropriate graphs/charts when explaining concepts
        - Differentiate between micro and macroeconomics concepts
        - Reference key economists and their theories when relevant
        - For calculations, show formulas and step-by-step working
        """,
        "examples": """
        **Example Q:** Explain price elasticity of demand?
        **Example A:** 
        **Price Elasticity of Demand** measures how quantity demanded responds to price changes...
        
        Formula: 
        ```
        PED = (%Δ Quantity Demanded) / (%Δ Price)
        ```
        
        Types:
        - Elastic (PED > 1)
        - Inelastic (PED < 1)
        - Unitary elastic (PED = 1)
        
        **Real-world example:** Gasoline often has inelastic demand...
        """,
        "question_types": {
            "Multiple Choice": "Include questions about economic models and graph interpretations",
            "True/False": "Focus on economic principles and common misconceptions",
            "Calculation": "Include problems with economic formulas and step-by-step solutions"
        }
    },
    "marketing": {
        "description": "Marketing strategies, consumer behavior, and campaign analysis",
        "guidance": """
        - Reference the 4Ps framework (Product, Price, Place, Promotion)
        - Use modern marketing examples when possible
        - Differentiate between B2B and B2C marketing
        - Include digital marketing considerations
        - Reference case studies where appropriate
        """,
        "examples": """
        **Example Q:** What makes an effective marketing campaign?
        **Example A:** Key elements include:
        1. **Clear target audience** definition
        2. **Value proposition** articulation
        3. Multi-channel **integration**
        4. Measurable **KPIs**
        
        **Example:** Apple's 'Think Different' campaign succeeded because...
        """,
        "question_types": {
            "Multiple Choice": "Include questions about marketing frameworks and case studies",
            "True/False": "Focus on marketing principles and consumer behavior",
            "Case Analysis": "Provide mini case studies for analysis"
        }
    },
    "finance": {
        "description": "Financial analysis, valuation, and investment strategies",
        "guidance": """
        - Use proper financial terminology and formulas
        - Show calculations step-by-step
        - Differentiate between corporate finance and investments
        - Reference key financial ratios and metrics
        - Use realistic numerical examples
        """,
        "examples": """
        **Example Q:** How do you calculate NPV?
        **Example A:**
        **Net Present Value (NPV)** calculates the present value of future cash flows:
        
        Formula:
        ```
        NPV = Σ [CFₜ / (1+r)ᵗ] - Initial Investment
        ```
        
        Where:
        - CFₜ = Cash flow at time t
        - r = Discount rate
        - t = Time period
        
        **Example:** For a project with $100 initial investment...
        """,
        "question_types": {
            "Multiple Choice": "Include questions about financial formulas and ratio analysis",
            "Calculation": "Provide financial problems requiring step-by-step solutions",
            "Scenario Analysis": "Present investment scenarios for evaluation"
        }
    },
    "general": {
        "description": "General academic subjects and course materials",
        "guidance": """
        - Provide clear, structured explanations
        - Use academic conventions appropriate for the material
        - Organize information logically
        - Highlight key terms and concepts
        """,
        "examples": "",
        "question_types": {
            "Multiple Choice": "Standard multiple choice questions",
            "True/False": "True/false statements",
            "Short Answer": "Brief open-ended questions"
        }
    }
}

# Create storage directory
PERSIST_DIR = "./storage"
Path(PERSIST_DIR).mkdir(exist_ok=True)

# Initialize session state
if 'processed_files' not in st.session_state:
    st.session_state.processed_files = False
if 'generated_mcqs' not in st.session_state:
    st.session_state.generated_mcqs = None
if 'last_embed_model' not in st.session_state:
    # Try to load from metadata if exists
    metadata_path = os.path.join(PERSIST_DIR, "metadata.json")
    if os.path.exists(metadata_path):
        try:
            with open(metadata_path, "r") as f:
                metadata = json.load(f)
            st.session_state.last_embed_model = metadata.get("embed_model", "text-embedding-3-small")
            st.session_state.last_subject = metadata.get("subject", "general")
        except:
            st.session_state.last_embed_model = "text-embedding-3-small"
            st.session_state.last_subject = "general"
    else:
        st.session_state.last_embed_model = "text-embedding-3-small"
        st.session_state.last_subject = "general"

# Sidebar for settings
with st.sidebar:
    st.header("⚙️ Settings")
    if st.button("🔄 Clear Index & Reload Documents", help="Wipe all stored data and start fresh"):
        if os.path.exists(PERSIST_DIR):
            shutil.rmtree(PERSIST_DIR)
        if os.path.exists("materials"):
            shutil.rmtree("materials")
        st.session_state.processed_files = False
        st.success("Cleared index and materials. Please upload your documents again.")
        time.sleep(2)
        st.rerun()
    
    # Get OpenAI API key
    openai_api_key = st.text_input(
        "🔑 OpenAI API Key",
        value="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",
        type="password",
        help="Get your API key from https://platform.openai.com/account/api-keys"
    )
    
    if openai_api_key:
        os.environ["OPENAI_API_KEY"] = openai_api_key
        openai.api_key = openai_api_key
    
    st.divider()
    st.markdown("### Teaching Subject")
    subject = st.selectbox(
        "Select Subject Specialization",
        list(SUBJECT_CONFIGS.keys()),
        index=list(SUBJECT_CONFIGS.keys()).index(st.session_state.get('last_subject', 'general')),
        help="Specializes the AI for specific academic disciplines"
    )
    st.caption(SUBJECT_CONFIGS[subject]['description'])
    
    st.divider()
    st.markdown("### Advanced Options")
    embed_model_name = st.selectbox(
        "Embedding Model",
        ["text-embedding-3-small", "text-embedding-3-large", "text-embedding-ada-002"],
        index=0,
        help="Model for converting text to vectors"
    )
    llm_model_name = st.selectbox(
        "LLM Model",
        ["gpt-3.5-turbo", "gpt-4", "gpt-4-turbo"],
        index=0,
        help="Language model for generating responses"
    )
    chunk_size = st.slider(
        "Chunk Size (characters)",
        256, 2048, 1024,
        help="Size of document chunks for processing"
    )

# File upload section
uploaded_files = st.file_uploader(
    "📤 Upload course materials (PDF, DOCX, TXT)",
    type=["pdf", "docx", "txt"],
    accept_multiple_files=True,
    help="Upload all relevant course materials at once"
)

def create_subject_specific_engine(index, subject, llm):
    config = SUBJECT_CONFIGS.get(subject.lower(), SUBJECT_CONFIGS['general'])
    
    # Build system prompt
    system_prompt = f"""
You are an expert {subject} teaching assistant helping with academic course materials. 
Your goal is to provide accurate, concise, and well-structured responses STRICTLY based on the provided materials.

SPECIALIZATION: {subject.upper()}
{config['guidance']}

RESPONSE RULES:
1. For concepts: Provide definitions, examples, and key characteristics
2. For procedures: List steps in numbered format
3. For comparisons: Use tables or bullet points
4. For calculations: Show formulas and step-by-step working
5. NEVER invent information not present in the documents
6. If unsure: State "This isn't covered in the provided materials"

FORMATTING:
- Use Markdown for clear formatting
- Highlight key terms in bold
- Use bullet points for lists
- Use tables for comparisons
- Include {subject}-specific examples

EXAMPLES:
{config['examples']}
"""
    
    return index.as_query_engine(
        similarity_top_k=6,
        response_mode="tree_summarize",
        streaming=False,
        verbose=True,
        llm=llm,
        system_prompt=system_prompt
    )

# Process files when uploaded
if uploaded_files and not st.session_state.processed_files:
    with st.spinner("Processing files..."):
        try:
            # Clear and recreate materials directory
            if os.path.exists("materials"):
                shutil.rmtree("materials")
            os.makedirs("materials")
            
            # Save uploaded files to materials directory
            for file in uploaded_files:
                with open(os.path.join("materials", file.name), "wb") as f:
                    f.write(file.getbuffer())
            
            # Initialize models with selected options
            embed_model = OpenAIEmbedding(
                model=embed_model_name,
                embed_batch_size=10
            )
            llm = OpenAI(
                model=llm_model_name,
                temperature=0.1,
                max_tokens=2000
            )
            
            Settings.llm = llm
            Settings.embed_model = embed_model
            Settings.chunk_size = chunk_size
            
            # Check if embedding model or subject has changed
            model_changed = (st.session_state.last_embed_model != embed_model_name or 
                           st.session_state.last_subject != subject)
            if model_changed:
                st.warning("Configuration changed! Rebuilding index...")
                if os.path.exists(PERSIST_DIR):
                    shutil.rmtree(PERSIST_DIR)
                    os.makedirs(PERSIST_DIR)
                st.session_state.last_embed_model = embed_model_name
                st.session_state.last_subject = subject
            
            # Load or create index
            if os.path.exists(PERSIST_DIR) and os.listdir(PERSIST_DIR) and not model_changed:
                storage_context = StorageContext.from_defaults(persist_dir=PERSIST_DIR)
                index = load_index_from_storage(storage_context)
                st.success("Loaded existing index from storage!")
            else:
                # Enhanced PDF reader with fallback                
                try:
                    # Use direct PDF extraction methods
                    from llama_index.core import SimpleDirectoryReader
                    from pdfminer.high_level import extract_text
                    from docx import Document as DocxDocument
                    import fitz  # PyMuPDF
                    from io import StringIO
                    
                    # Custom PDF loader that uses both PyMuPDF and pdfminer for maximum compatibility
                    def load_pdf_with_fallback(file_path):
                        file_name = os.path.basename(file_path)
                        try:
                            # Try PyMuPDF first (usually better quality)
                            doc = fitz.open(file_path)
                            text = ""
                            for page_num in range(len(doc)):
                                page = doc.load_page(page_num)
                                text += page.get_text() + "\n\n"
                            
                            if len(text.strip()) > 100:
                                return [Document(text=text, metadata={
                                    "file_name": file_name,
                                    "page_count": len(doc)
                                })]
                        except Exception as e:
                            st.warning(f"PyMuPDF failed for {file_name}, trying pdfminer: {str(e)}")
                        
                        # Fallback to pdfminer
                        try:
                            text = extract_text(file_path)
                            if len(text.strip()) > 50:
                                return [Document(text=text, metadata={"file_name": file_name})]
                        except Exception as e:
                            st.error(f"PDF extraction failed for {file_name}: {str(e)}")
                        
                        return []
                    
                    # Configure file extractors
                    file_extractor = {
                        ".pdf": load_pdf_with_fallback,
                        ".docx": lambda f: [Document(text="\n".join([p.text for p in DocxDocument(f).paragraphs]), 
                                           metadata={"file_name": os.path.basename(f)})],
                        ".txt": lambda f: [Document(text=open(f, 'r', encoding='utf-8').read(), 
                                          metadata={"file_name": os.path.basename(f)})]
                    }
                    
                    documents = []
                    for file_path in Path("materials").rglob("*"):
                        if file_path.suffix.lower() in file_extractor:
                            docs = file_extractor[file_path.suffix.lower()](str(file_path))
                            if docs:
                                documents.extend(docs)
                    
                except ImportError as e:
                    st.error(f"Required libraries not installed: {str(e)}")
                    st.info("Run: pip install pymupdf pdfminer.six python-docx")
                    st.stop()
                
                # Validate documents
                valid_docs = []
                for doc in documents:
                    clean_text = ' '.join(doc.text.strip().split())
                    if len(clean_text) > 50:  # Minimum viable content length
                        # Create a new document with the cleaned text instead of modifying existing one
                        new_doc = Document(text=clean_text, metadata=doc.metadata)
                        valid_docs.append(new_doc)
                
                if not valid_docs:
                    st.error("No readable content found in uploaded files.")
                    st.stop()
                
                # Debug output in expander
                with st.expander("🔍 Document Processing Details", expanded=False):
                    st.write(f"### Processed {len(valid_docs)} documents")
                    for i, doc in enumerate(valid_docs[:3]):  # Show first 3 max
                        st.write(f"#### Document {i+1}: {os.path.basename(doc.metadata.get('file_name', 'unknown'))}")
                        
                        # Clean metadata display
                        clean_metadata = {k:v for k,v in doc.metadata.items() 
                                        if not isinstance(v, (bytes, bytearray))}
                        st.json(clean_metadata)
                        
                        # Text preview
                        st.write("**Content preview:**")
                        st.text(doc.text[:500] + ("..." if len(doc.text) > 500 else ""))
                        st.divider()
                
                # Create index
                index = VectorStoreIndex.from_documents(
                    valid_docs,
                    show_progress=True,
                    embed_model=embed_model
                )
                
                # Save index with metadata
                index.storage_context.persist(persist_dir=PERSIST_DIR)
                
                # Store metadata
                metadata = {
                    "embed_model": embed_model_name,
                    "llm_model": llm_model_name,
                    "chunk_size": chunk_size,
                    "subject": subject,
                    "file_count": len(valid_docs),
                    "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
                }
                with open(os.path.join(PERSIST_DIR, "metadata.json"), "w") as f:
                    json.dump(metadata, f)
                
                st.success(f"Processed {len(valid_docs)} documents successfully!")
            
            # Create subject-specific query engine
            query_engine = create_subject_specific_engine(index, subject, llm)
            
            st.session_state.query_engine = query_engine
            st.session_state.processed_files = True
            st.session_state.current_subject = subject
            
        except Exception as e:
            st.error(f"Error processing files: {str(e)}")
            st.stop()

# Main interface tabs
if st.session_state.get('processed_files', False):
    tab1, tab2, tab3 = st.tabs(["💬 Ask Questions", "📝 Generate Practice Questions", "🔑 Key Concepts"])
    
    with tab1:
        st.header(f"💬 Ask {st.session_state.current_subject.capitalize()} Questions")
        user_question = st.text_area(
            "Enter your question:",
            placeholder=f"e.g. {SUBJECT_CONFIGS[subject]['examples'].split('?')[0][8:]+'?' if SUBJECT_CONFIGS[subject]['examples'] else 'Explain a key concept from the materials...'}",
            height=100
        )
        
        if user_question:
            with st.spinner("Analyzing materials..."):
                try:
                    start_time = time.time()
                    response = st.session_state.query_engine.query(user_question)
                    end_time = time.time()
                    
                    st.markdown("### 📝 Answer")
                    st.write(str(response))
                    
                    # Enhanced source display
                    if hasattr(response, 'source_nodes') and response.source_nodes:
                        st.markdown("---")
                        st.markdown("#### 📚 Source References")
                        
                        unique_sources = {}
                        for node in response.source_nodes:
                            source_text = node.node.text
                            file_name = node.node.metadata.get('file_name', 'Unknown')
                            page_num = node.node.metadata.get('page_label', '')
                            
                            if file_name not in unique_sources:
                                unique_sources[file_name] = {
                                    'text': source_text,
                                    'pages': {page_num} if page_num else set()
                                }
                            elif page_num:
                                unique_sources[file_name]['pages'].add(page_num)
                        
                        for file_name, data in unique_sources.items():
                            pages = ", ".join(sorted(data['pages'])) if data['pages'] else "N/A"
                            st.write(f"**📄 {file_name}** (Pages: {pages})")
                            st.caption(f"*Relevant excerpt:* {data['text'][:200]}...")
                    
                    st.info(f"⏱️ Generated in {end_time - start_time:.2f} seconds")
                
                except Exception as e:
                    st.error(f"Error generating answer: {str(e)}")
                    st.info("Try rephrasing your question or check the document content.")
    
    with tab2:
        st.header("📝 Generate Practice Questions")
        
        col1, col2 = st.columns(2)
        with col1:
            topic = st.text_input(
                "Topic for questions:",
                placeholder=f"e.g. {'Supply and demand' if subject == 'economics' else 'Marketing mix' if subject == 'marketing' else 'Financial ratios'}",
                key="q_topic"
            )
            
            # Get subject-specific question types
            q_types = list(SUBJECT_CONFIGS[subject]['question_types'].keys())
            question_type = st.selectbox(
                "Question type",
                q_types,
                key="q_type"
            )
            st.caption(SUBJECT_CONFIGS[subject]['question_types'][question_type])
        
        with col2:
            difficulty = st.select_slider(
                "Difficulty level",
                options=["Easy", "Medium", "Hard"],
                value="Medium",
                key="q_difficulty"
            )
            num_questions = st.slider(
                "Number of questions",
                1, 15, 5,
                key="q_count"
            )
        
        if st.button("Generate Questions", key="generate_q") and topic:
            with st.spinner(f"Generating {num_questions} {question_type.lower()} questions..."):
                try:
                    # Get relevant context
                    context_prompt = f"""
Extract comprehensive information about '{topic}' from the course materials including:
- Key definitions
- Important examples
- Relevant formulas/theories
- Practical applications
"""
                    context_response = st.session_state.query_engine.query(context_prompt)
                    relevant_content = str(context_response)
                    
                    # Generate questions based on type
                    if question_type == "Multiple Choice":
                        prompt = f"""
Generate {num_questions} {difficulty.lower()} multiple choice questions about '{topic}' for {subject}.
Base questions strictly on this content: {relevant_content}

Format each question EXACTLY like this:

### Question [N]

**[Q].** [Question text]

- **A)** [Option A]
- **B)** [Option B]
- **C)** [Option C]
- **D)** [Option D]

**Answer:** [Correct letter]
**Explanation:** [1-2 sentence explanation with reference to {subject} concepts]

---
"""
                    elif question_type == "True/False":
                        prompt = f"""
Generate {num_questions} true/false questions about '{topic}' for {subject}.
Include explanations and reference specific content.

Format EXACTLY like this:

### Question [N]

**[Q].** [Statement]

**Answer:** True/False
**Explanation:** [Explanation with reference to {subject} materials]

---
"""
                    elif question_type in ["Calculation", "Case Analysis"]:
                        prompt = f"""
Generate {num_questions} {question_type.lower()} problems about '{topic}' for {subject}.
Include detailed solutions and explanations.

Format EXACTLY like this:

### Problem [N]

**[Q].** [Problem statement]

**Solution:** 
[Step-by-step solution with {subject}-specific reasoning]

**Key Takeaway:**
[Main learning point]

---
"""
                    else:  # Short Answer or Essay
                        prompt = f"""
Generate {num_questions} {question_type.lower()} questions about '{topic}' for {subject}.
Include detailed answers and references.

Format EXACTLY like this:

### Question [N]

**[Q].** [Question]

**Answer:** 
[Detailed answer with references to {subject} materials]

---
"""
                    
                    # Generate questions
                    response = openai.chat.completions.create(
                        model=llm_model_name,
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.3,
                        max_tokens=2000
                    )
                    
                    questions = response.choices[0].message.content
                    st.session_state.generated_mcqs = questions
                    
                    st.markdown("### Generated Questions")
                    st.markdown(questions)
                    
                    # Download button
                    st.download_button(
                        label="📥 Download Questions",
                        data=questions,
                        file_name=f"{subject}_{topic.replace(' ', '_')}_questions.md",
                        mime="text/markdown"
                    )
                
                except Exception as e:
                    st.error(f"Error generating questions: {str(e)}")
    
    with tab3:
        st.header("🔑 Extract Key Concepts")
        concept_topic = st.text_input(
            "Enter topic (leave blank for entire document):",
            placeholder=f"e.g. {'Macroeconomics' if subject == 'economics' else 'Consumer behavior' if subject == 'marketing' else 'Valuation methods'}",
            key="concept_topic"
        )
        
        if st.button("Extract Key Points", key="extract_concepts"):
            with st.spinner("Identifying key concepts..."):
                try:
                    if concept_topic:
                        prompt = f"""
Extract the 7-10 most important {subject} concepts about '{concept_topic}' from the materials.
For each concept include:
1. Definition
2. Key characteristics
3. {subject}-specific examples
4. Practical applications

Format as:

### Key {subject.capitalize()} Concepts: {concept_topic}

1. **[Concept Name]**: [Brief definition/explanation]
   - **Key Features**: [List characteristics]
   - **Example**: [Relevant example]
   - **Application**: [How it's used in {subject}]

2. **[Concept Name]**: [Brief definition/explanation]
   ...
"""
                    else:
                        prompt = f"""
Extract the 10-15 most important {subject} concepts from the entire document.
Organize by major topics and include {subject}-specific context.

Format as:

### Document Key {subject.capitalize()} Concepts

#### [Major Topic 1]
1. [Concept] - [Explanation with {subject} context]
2. [Concept] - [Explanation with {subject} context]

#### [Major Topic 2]
...
"""
                    
                    response = st.session_state.query_engine.query(prompt)
                    concepts = str(response)
                    
                    st.markdown("### 📌 Key Concepts")
                    st.markdown(concepts)
                    
                    st.download_button(
                        label="📥 Download Concepts",
                        data=concepts,
                        file_name=f"{subject}_key_concepts.md",
                        mime="text/markdown"
                    )
                
                except Exception as e:
                    st.error(f"Error extracting concepts: {str(e)}")

elif uploaded_files and not st.session_state.processed_files:
    st.warning("⏳ Files uploaded but not yet processed. Please wait...")

else:
    st.info("📤 Please upload course materials to begin.")
    st.image("https://via.placeholder.com/800x400?text=Upload+PDF%2FDOCX%2FTXT+Files", use_column_width=True)

# Footer
st.markdown("---")
st.caption("AI Teaching Assistant v3.0 | Subject-Specialized | For educational purposes only")