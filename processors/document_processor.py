"""
Document processing utilities
"""

import os
import json
import time
import streamlit as st
from pathlib import Path
from llama_index.core import (
    VectorStoreIndex,
    Document
)
from config.settings import PERSIST_DIR


def load_pdf_with_fallback(file_path):
    """Load PDF with fallback mechanism for better extraction"""
    try:
        from pdfminer.high_level import extract_text
        import fitz  # PyMuPDF
        
        file_name = os.path.basename(file_path)
        try:
            doc = fitz.open(file_path)
            text = ""
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text() + "\n\n"
            
            if len(text.strip()) > 100:
                return [Document(text=text, metadata={
                    "file_name": file_name,
                    "page_count": len(doc)
                })]
        except Exception as e:
            st.warning(f"PyMuPDF a échoué pour {file_name}, essai avec pdfminer: {str(e)}")
        
        try:
            text = extract_text(file_path)
            if len(text.strip()) > 50:
                return [Document(text=text, metadata={"file_name": file_name})]
        except Exception as e:
            st.error(f"L'extraction PDF a échoué pour {file_name}: {str(e)}")
    except ImportError:
        st.error("Bibliothèques requises non installées: PyMuPDF ou pdfminer.six")
        st.info("Exécutez: pip install pymupdf pdfminer.six python-docx")
    
    return []


def process_documents(paths, embed_model):
    """Process documents from file paths"""
    try:
        from docx import Document as DocxDocument
        
        # Configure file extractors
        file_extractor = {
            ".pdf": load_pdf_with_fallback,
            ".docx": lambda f: [Document(text="\n".join([p.text for p in DocxDocument(f).paragraphs]), 
                               metadata={"file_name": os.path.basename(f)})],
            ".txt": lambda f: [Document(text=open(f, 'r', encoding='utf-8').read(), 
                              metadata={"file_name": os.path.basename(f)})]
        }
        
        documents = []
        for file_path in paths:
            suffix = Path(file_path).suffix.lower()
            if suffix in file_extractor:
                docs = file_extractor[suffix](str(file_path))
                if docs:
                    documents.extend(docs)
        
        # Validate documents
        valid_docs = []
        for doc in documents:
            clean_text = ' '.join(doc.text.strip().split())
            if len(clean_text) > 50:
                new_doc = Document(text=clean_text, metadata=doc.metadata)
                valid_docs.append(new_doc)
        
        if not valid_docs:
            return None
        
        # Create index
        index = VectorStoreIndex.from_documents(
            valid_docs,
            show_progress=True,
            embed_model=embed_model
        )
        
        return index, valid_docs
        
    except ImportError as e:
        st.error(f"Bibliothèques requises non installées: {str(e)}")
        st.info("Exécutez: pip install pymupdf pdfminer.six python-docx")
        return None, []


def save_index(index, embed_model_name, llm_model_name, chunk_size, subject, language, valid_docs):
    """Save index and metadata"""
    # Save index
    index.storage_context.persist(persist_dir=PERSIST_DIR)
    
    # Store metadata
    metadata = {
        "embed_model": embed_model_name,
        "llm_model": llm_model_name,
        "chunk_size": chunk_size,
        "subject": subject,
        "language": language,
        "file_count": len(valid_docs),
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
    }
    
    with open(os.path.join(PERSIST_DIR, "metadata.json"), "w") as f:
        json.dump(metadata, f)
    
    return True
